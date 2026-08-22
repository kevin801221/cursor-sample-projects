#!/usr/bin/env python3
"""
Step 0: 統一入口——YouTube / PDF / DOCX / 網頁 URL 都收
==========================================================
把任何來源正規化成同一個 source.json 格式，下游 Phase 2~5 完全不用改。

驗證環境: Python 3.12 / pymupdf 1.28.2 / python-docx 1.2.0 / trafilatura 2.2.0

統一輸出格式（與 01 的 transcript.json 相容）:
{
  "video_id": "<source_id>",      # 沿用欄位名，下游冪等/圖譜都靠它
  "source_type": "youtube|pdf|docx|url",
  "title": ..., "url": ...,
  "segments": [{"start": <秒或頁碼或段序>, "end": ..., "text": ...,
                "ref": "<可回溯的引用位置：時間戳連結/檔案#page=N/網址>"}]
}

付費 vs 免費雙路線（教學點：agent 執行前要問使用者選哪條）:
  PDF : 免費地端 pymupdf（本腳本內建，純文字 PDF 效果好）
        付費雲端 LlamaParse（--engine llamaparse，表格/掃描件效果好，
        需 LLAMA_CLOUD_API_KEY）
  URL : 免費地端 trafilatura（本腳本內建，直接抓網頁正文）
        付費 Tavily Extract（--engine tavily，反爬/動態頁面成功率高，
        需 TAVILY_API_KEY）
  DOCX: python-docx 免費地端即可，無付費必要（教學點：不是所有環節
        都值得花錢，工具選擇要看邊際效益）

用法:
  python 00_ingest_source.py "https://www.youtube.com/watch?v=..." --out source.json
  python 00_ingest_source.py report.pdf --out source.json
  python 00_ingest_source.py report.pdf --engine llamaparse --out source.json
  python 00_ingest_source.py spec.docx --out source.json
  python 00_ingest_source.py "https://blog.example.com/post" --out source.json
  python 00_ingest_source.py "https://blog.example.com/post" --engine tavily --out source.json
"""
import argparse
import hashlib
import json
import os
import re
import sys
from pathlib import Path

MIN_SEG_CHARS = 200  # 文件段落聚合下限：太碎的段落合併到這個量級再輸出


def _source_id(seed: str) -> str:
    return hashlib.sha1(seed.encode("utf-8")).hexdigest()[:12]


def _pack_paragraphs(paras: list[tuple[int, str]], ref_fn) -> list[dict]:
    """把 (位置, 文字) 列表聚合成 >= MIN_SEG_CHARS 的段。ref 取該段起始位置。"""
    segments, buf, buf_start = [], [], None
    for pos, text in paras:
        text = text.strip()
        if not text:
            continue
        if buf_start is None:
            buf_start = pos
        buf.append(text)
        if sum(len(t) for t in buf) >= MIN_SEG_CHARS:
            segments.append({"start": buf_start, "end": pos,
                             "text": "\n".join(buf), "ref": ref_fn(buf_start)})
            buf, buf_start = [], None
    if buf:
        segments.append({"start": buf_start, "end": paras[-1][0] if paras else buf_start,
                         "text": "\n".join(buf), "ref": ref_fn(buf_start)})
    return segments


# ---------------- PDF ----------------
def ingest_pdf_local(path: Path) -> dict:
    """免費地端: pymupdf 逐頁抽文字。start/end = 頁碼(1-based)，ref 帶 #page=N。"""
    import pymupdf

    doc = pymupdf.open(path)
    paras = [(i + 1, page.get_text("text")) for i, page in enumerate(doc)]
    doc.close()
    abs_ref = path.resolve().as_uri()
    segments = _pack_paragraphs(paras, lambda p: f"{abs_ref}#page={p}")
    if not segments:
        print("[!] 抽不到文字——可能是掃描版 PDF。請改用 --engine llamaparse（OCR）",
              file=sys.stderr)
        sys.exit(2)
    return {"video_id": _source_id(str(path.resolve())), "source_type": "pdf",
            "title": path.stem, "url": abs_ref, "segments": segments}


def ingest_pdf_llamaparse(path: Path) -> dict:
    """付費雲端: LlamaParse。需 `uv sync --extra paid` 及 LLAMA_CLOUD_API_KEY。"""
    if not os.environ.get("LLAMA_CLOUD_API_KEY"):
        print("[!] 缺 LLAMA_CLOUD_API_KEY（https://cloud.llamaindex.ai 免費註冊有額度）",
              file=sys.stderr)
        sys.exit(2)
    from llama_parse import LlamaParse  # 付費路線才需要，所以延遲 import

    parser = LlamaParse(result_type="markdown")
    docs = parser.load_data(str(path))
    abs_ref = path.resolve().as_uri()
    paras = [(i + 1, d.text) for i, d in enumerate(docs)]
    segments = _pack_paragraphs(paras, lambda p: f"{abs_ref}#page={p}")
    return {"video_id": _source_id(str(path.resolve())), "source_type": "pdf",
            "title": path.stem, "url": abs_ref, "segments": segments}


# ---------------- DOCX ----------------
def ingest_docx(path: Path) -> dict:
    """免費地端: python-docx。start/end = 段落序號，含表格文字。"""
    import docx

    d = docx.Document(str(path))
    paras = [(i + 1, p.text) for i, p in enumerate(d.paragraphs)]
    base = len(paras)
    for ti, table in enumerate(d.tables):  # 表格常藏關鍵資料，別漏
        rows = [" | ".join(c.text.strip() for c in row.cells) for row in table.rows]
        paras.append((base + ti + 1, "\n".join(r for r in rows if r.strip(" |"))))
    abs_ref = path.resolve().as_uri()
    segments = _pack_paragraphs(paras, lambda p: f"{abs_ref}#para={p}")
    if not segments:
        print("[!] DOCX 內沒有可抽取文字", file=sys.stderr)
        sys.exit(2)
    return {"video_id": _source_id(str(path.resolve())), "source_type": "docx",
            "title": path.stem, "url": abs_ref, "segments": segments}


# ---------------- URL ----------------
def extract_html(html: str | bytes, url: str) -> tuple[str | None, str | None]:
    """trafilatura 正文抽取（獨立函式，方便離線測試）。回傳 (title, text)。"""
    import trafilatura

    text = trafilatura.extract(html, url=url, include_comments=False)
    meta = trafilatura.extract_metadata(html)
    title = meta.title if meta else None
    return title, text


def ingest_url_free(url: str) -> dict:
    """免費地端: requests + trafilatura。"""
    import requests

    resp = requests.get(url, timeout=30,
                        headers={"User-Agent": "Mozilla/5.0 (teaching-bot)"})
    resp.raise_for_status()
    # 餵 bytes 讓 trafilatura 自己偵測編碼：header 沒宣告 charset 時
    # resp.text 會退回 latin-1，中文全變亂碼
    title, text = extract_html(resp.content, url)
    if not text:
        print("[!] 正文抽取失敗（可能是 JS 渲染頁）。改用 --engine tavily 成功率較高",
              file=sys.stderr)
        sys.exit(2)
    paras = [(i + 1, p) for i, p in enumerate(re.split(r"\n{2,}", text))]
    segments = _pack_paragraphs(paras, lambda p: url)
    return {"video_id": _source_id(url), "source_type": "url",
            "title": title or url, "url": url, "segments": segments}


def ingest_url_tavily(url: str) -> dict:
    """付費: Tavily Extract API。需 TAVILY_API_KEY（tavily.com 免費註冊含額度）。"""
    if not os.environ.get("TAVILY_API_KEY"):
        print("[!] 缺 TAVILY_API_KEY（https://tavily.com 免費註冊每月有免費額度）",
              file=sys.stderr)
        sys.exit(2)
    import requests

    resp = requests.post("https://api.tavily.com/extract",
                         json={"api_key": os.environ["TAVILY_API_KEY"], "urls": [url]},
                         timeout=60)
    resp.raise_for_status()
    results = resp.json().get("results", [])
    if not results or not results[0].get("raw_content"):
        print("[!] Tavily 也抽不到內容，請確認網址可公開存取", file=sys.stderr)
        sys.exit(2)
    text = results[0]["raw_content"]
    paras = [(i + 1, p) for i, p in enumerate(re.split(r"\n{2,}", text))]
    segments = _pack_paragraphs(paras, lambda p: url)
    return {"video_id": _source_id(url), "source_type": "url",
            "title": url, "url": url, "segments": segments}


# ---------------- Router ----------------
YT_PAT = re.compile(r"(youtube\.com/watch|youtu\.be/|youtube\.com/shorts)")


def main(source: str, engine: str | None, out: Path, langs: list[str]):
    if YT_PAT.search(source):
        # YouTube 走 Step 1 既有邏輯（含 CC/自動字幕/whisper fallback）
        import importlib.util
        spec = importlib.util.spec_from_file_location(
            "yt", Path(__file__).parent / "01_fetch_transcript.py")
        yt = importlib.util.module_from_spec(spec)
        spec.loader.exec_module(yt)
        result = yt.fetch(source, langs, out)
        # 補統一欄位：segments 加 ref（時間戳深連結）
        for seg in result["segments"]:
            seg["ref"] = f"https://youtu.be/{result['video_id']}?t={int(seg['start'])}"
        result["source_type"] = "youtube"
        out.write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8")
        # 輸出格式與 pdf/docx/url 路線一致，SKILL.md Phase 1 的成功判準才對得上
        print(f"[✓] {result['source_type']} 來源 {len(result['segments'])} 段已存至 {out}")
        return

    if source.startswith(("http://", "https://")):
        result = ingest_url_tavily(source) if engine == "tavily" else ingest_url_free(source)
    else:
        p = Path(source)
        if not p.exists():
            print(f"[!] 找不到檔案: {p}", file=sys.stderr)
            sys.exit(1)
        suffix = p.suffix.lower()
        if suffix == ".pdf":
            result = ingest_pdf_llamaparse(p) if engine == "llamaparse" else ingest_pdf_local(p)
        elif suffix == ".docx":
            result = ingest_docx(p)
        else:
            print(f"[!] 不支援的格式: {suffix}（支援 .pdf/.docx/YouTube/網頁 URL）",
                  file=sys.stderr)
            sys.exit(1)

    out.write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"[✓] {result['source_type']} 來源 {len(result['segments'])} 段已存至 {out}")


if __name__ == "__main__":
    ap = argparse.ArgumentParser()
    ap.add_argument("source", help="YouTube URL / 檔案路徑(.pdf/.docx) / 網頁 URL")
    ap.add_argument("--engine", choices=["llamaparse", "tavily"],
                    help="付費引擎（預設用免費地端工具）")
    ap.add_argument("--out", type=Path, default=Path("source.json"))
    ap.add_argument("--lang", nargs="+", default=["zh-TW", "zh-Hant", "zh", "en"],
                    help="YouTube 字幕語言偏好")
    args = ap.parse_args()
    main(args.source, args.engine, args.out, args.lang)
