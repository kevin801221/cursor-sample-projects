#!/usr/bin/env python3
"""
ingest_pipeline.py
==================
支援 YouTube URL、PDF、DOCX、網頁 URL 的端到端 GraphRAG 入庫管線：
1. 來源擷取與正規化 (YouTube yt-dlp 字幕 / PDF PyMuPDF / DOCX python-docx / URL trafilatura)
2. 語意聚合切塊 & Chroma 向量資料庫嵌入
3. LLM (Gemini) 知識圖譜三元組抽取 & Neo4j 入庫
"""
import hashlib
import json
import os
import re
import shutil
import tempfile
from pathlib import Path
from typing import Callable, Optional

import yt_dlp
from langchain_chroma import Chroma
from langchain_core.documents import Document
from langchain_google_genai import ChatGoogleGenerativeAI, GoogleGenerativeAIEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from neo4j import GraphDatabase

MIN_SEG_CHARS = 200
WINDOW_SECONDS = 60
CHUNK_SIZE = 800
CHUNK_OVERLAP = 120

CHAT_MODEL = os.environ.get("GEMINI_MODEL", "gemini-3.5-flash")
EMBED_MODEL = os.environ.get("GEMINI_EMBED_MODEL", "gemini-embedding-001")
EXTRACT_PROMPT = """你是知識圖譜抽取器。從以下內容片段抽取三元組。
規則:
- entity 用正式名稱（如 "LangGraph"、"注意力機制"、"Transformer"），同義詞正規化為同一名稱
- relation 用簡短動詞片語（如 "包含"、"用於"、"解決"、"依賴於"）
- 只抽取片段中明確陳述的關係，不要腦補
- 最多 8 個三元組
- 嚴格只輸出 JSON，不要 markdown 圍欄，格式:
{"triples": [{"subject": "...", "relation": "...", "object": "..."}]}

片段:
"""

YT_PAT = re.compile(r"(youtube\.com/watch|youtu\.be/|youtube\.com/shorts)")


def _source_id(seed: str) -> str:
    return hashlib.sha1(seed.encode("utf-8")).hexdigest()[:12]


def _file_id(path: Path) -> str:
    # 用檔案內容而非路徑做 id：同一份檔案不管從哪裡（含上傳暫存路徑）進來，
    # 都會對上同一筆資料，冪等清舊才有效
    return hashlib.sha1(path.read_bytes()).hexdigest()[:12]


def _pack_paragraphs(paras: list[tuple[int, str]], ref_fn: Callable[[int], str]) -> list[dict]:
    segments, buf, buf_start = [], [], None
    for pos, text in paras:
        text = text.strip()
        if not text:
            continue
        if buf_start is None:
            buf_start = pos
        buf.append(text)
        if sum(len(t) for t in buf) >= MIN_SEG_CHARS:
            segments.append({
                "start": buf_start,
                "end": pos,
                "text": "\n".join(buf),
                "ref": ref_fn(buf_start)
            })
            buf, buf_start = [], None
    if buf:
        segments.append({
            "start": buf_start,
            "end": paras[-1][0] if paras else buf_start,
            "text": "\n".join(buf),
            "ref": ref_fn(buf_start)
        })
    return segments


def pick_subtitle_lang(available: dict, preferred: list[str]) -> Optional[str]:
    if not available:
        return None
    for want in preferred:
        for lang in available:
            if lang == want or lang.startswith(want + "-") or want.startswith(lang + "-"):
                return lang
    return next(iter(available))


def parse_vtt(vtt_text: str) -> list[dict]:
    ts = re.compile(
        r"(\d{2}):(\d{2}):(\d{2})\.(\d{3})\s*-->\s*(\d{2}):(\d{2}):(\d{2})\.(\d{3})"
    )
    tag = re.compile(r"<[^>]+>")
    segments: list[dict] = []
    cur_start = cur_end = None
    buf: list[str] = []

    def flush():
        nonlocal cur_start, cur_end, buf
        if cur_start is None:
            return
        text = " ".join(t for t in buf if t).strip()
        text = re.sub(r"\s+", " ", text)
        if text:
            if segments and (text in segments[-1]["text"] or segments[-1]["text"].endswith(text)):
                segments[-1]["end"] = cur_end
            elif segments and text.startswith(segments[-1]["text"]):
                extra = text[len(segments[-1]["text"]):].strip()
                if extra:
                    segments.append({"start": cur_start, "end": cur_end, "text": extra})
                else:
                    segments[-1]["end"] = cur_end
            else:
                segments.append({"start": cur_start, "end": cur_end, "text": text})
        cur_start = cur_end = None
        buf = []

    for raw in vtt_text.splitlines():
        line = raw.strip()
        m = ts.search(line)
        if m:
            flush()
            h1, m1, s1, ms1, h2, m2, s2, ms2 = map(int, m.groups())
            cur_start = h1 * 3600 + m1 * 60 + s1 + ms1 / 1000
            cur_end = h2 * 3600 + m2 * 60 + s2 + ms2 / 1000
            continue
        if not line or line == "WEBVTT" or line.startswith(("Kind:", "Language:", "NOTE", "STYLE")):
            continue
        if cur_start is not None:
            buf.append(tag.sub("", line).strip())
    flush()
    return segments


# ---------------- 來源擷取 ----------------

def ingest_youtube(url: str, preferred_langs: list[str] = None) -> dict:
    if preferred_langs is None:
        preferred_langs = ["zh-TW", "zh-Hant", "zh", "en"]
    probe_opts = {"skip_download": True, "quiet": True, "no_warnings": True}
    with yt_dlp.YoutubeDL(probe_opts) as ydl:
        info = ydl.extract_info(url, download=False)

    manual = info.get("subtitles") or {}
    auto = info.get("automatic_captions") or {}

    source, lang = None, None
    if (lang := pick_subtitle_lang(manual, preferred_langs)):
        source = "manual"
    elif (lang := pick_subtitle_lang(auto, preferred_langs)):
        source = "auto"

    if source is None:
        raise ValueError(
            f"該影片沒有可用字幕 (url: {url})。"
            "可改走地端轉錄：yt-dlp -x 下載音訊後用 faster-whisper 轉成文字再匯入"
        )

    with tempfile.TemporaryDirectory() as tmpdir:
        dl_opts = {
            "skip_download": True,
            "writesubtitles": source == "manual",
            "writeautomaticsub": source == "auto",
            "subtitleslangs": [lang],
            "subtitlesformat": "vtt",
            "outtmpl": os.path.join(tmpdir, "%(id)s"),
            "quiet": True,
            "no_warnings": True,
        }
        with yt_dlp.YoutubeDL(dl_opts) as ydl:
            ydl.download([url])

        vtt_files = (list(Path(tmpdir).glob(f"{info['id']}.{lang}.vtt"))
                     or list(Path(tmpdir).glob(f"{info['id']}*.vtt")))
        if not vtt_files:
            raise RuntimeError("字幕下載失敗，請稍後重試")

        segments = parse_vtt(vtt_files[0].read_text(encoding="utf-8"))

    for seg in segments:
        seg["ref"] = f"https://youtu.be/{info['id']}?t={int(seg['start'])}"

    return {
        "video_id": info["id"],
        "source_type": "youtube",
        "title": info.get("title") or "YouTube 影片",
        "channel": info.get("channel"),
        "duration": info.get("duration"),
        "url": url,
        "subtitle_source": source,
        "subtitle_lang": lang,
        "segments": segments,
    }


def ingest_pdf_local(path: Path) -> dict:
    import pymupdf

    doc = pymupdf.open(path)
    paras = [(i + 1, page.get_text("text")) for i, page in enumerate(doc)]
    doc.close()
    abs_ref = path.resolve().as_uri()
    segments = _pack_paragraphs(paras, lambda p: f"{abs_ref}#page={p}")
    if not segments:
        raise ValueError("PDF 內沒有可抽取的文字，可能是掃描件（可改用 --engine llamaparse）")
    return {
        "video_id": _file_id(path),
        "source_type": "pdf",
        "title": path.stem,
        "url": abs_ref,
        "segments": segments,
    }


def ingest_pdf_llamaparse(path: Path) -> dict:
    if not os.environ.get("LLAMA_CLOUD_API_KEY"):
        raise ValueError("缺少 LLAMA_CLOUD_API_KEY")
    from llama_parse import LlamaParse

    parser = LlamaParse(result_type="markdown")
    docs = parser.load_data(str(path))
    abs_ref = path.resolve().as_uri()
    paras = [(i + 1, d.text) for i, d in enumerate(docs)]
    segments = _pack_paragraphs(paras, lambda p: f"{abs_ref}#page={p}")
    return {
        "video_id": _file_id(path),
        "source_type": "pdf",
        "title": path.stem,
        "url": abs_ref,
        "segments": segments,
    }


def ingest_docx(path: Path) -> dict:
    import docx

    d = docx.Document(str(path))
    paras = [(i + 1, p.text) for i, p in enumerate(d.paragraphs)]
    base = len(paras)
    for ti, table in enumerate(d.tables):
        rows = [" | ".join(c.text.strip() for c in row.cells) for row in table.rows]
        paras.append((base + ti + 1, "\n".join(r for r in rows if r.strip(" |"))))
    abs_ref = path.resolve().as_uri()
    segments = _pack_paragraphs(paras, lambda p: f"{abs_ref}#para={p}")
    if not segments:
        raise ValueError("DOCX 內沒有可抽取文字")
    return {
        "video_id": _file_id(path),
        "source_type": "docx",
        "title": path.stem,
        "url": abs_ref,
        "segments": segments,
    }


def _assert_public_url(url: str):
    # SSRF 防護：/ingest 可被瀏覽器跨域打到，不擋內網位址的話，
    # 惡意網頁能讓 server 去抓 metadata service / 內網服務再經 /chat 外洩。
    # ponytail: 只在發請求前驗一次，不防 DNS rebinding；
    # 測試要灌 localhost 頁面時設 ALLOW_PRIVATE_URLS=1
    if os.environ.get("ALLOW_PRIVATE_URLS"):
        return
    import ipaddress
    import socket
    from urllib.parse import urlparse

    host = urlparse(url).hostname
    if not host:
        raise ValueError(f"無法解析網址: {url}")
    try:
        infos = socket.getaddrinfo(host, None)
    except socket.gaierror:
        raise ValueError(f"無法解析主機名稱: {host}")
    for info in infos:
        ip = ipaddress.ip_address(info[4][0])
        if not ip.is_global:
            raise ValueError(
                f"拒絕存取內網/保留位址: {host} ({ip})。測試環境可設 ALLOW_PRIVATE_URLS=1"
            )


def ingest_url_free(url: str) -> dict:
    import requests
    import trafilatura

    headers = {"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64)"}
    # 手動逐跳跟隨重導向：3xx 轉向內網位址會繞過只驗第一個網址的防護
    for _ in range(5):
        _assert_public_url(url)
        resp = requests.get(url, timeout=30, headers=headers, allow_redirects=False)
        if resp.is_redirect or resp.is_permanent_redirect:
            url = requests.compat.urljoin(url, resp.headers["Location"])
            continue
        break
    else:
        raise ValueError("重導向次數過多")
    resp.raise_for_status()
    # 餵 bytes 讓 trafilatura 自己偵測編碼：header 沒宣告 charset 時
    # resp.text 會退回 latin-1，中文全變亂碼
    text = trafilatura.extract(resp.content, url=url, include_comments=False)
    meta = trafilatura.extract_metadata(resp.content)
    title = meta.title if meta else url
    if not text:
        raise ValueError("網頁正文抽取失敗，可能為動態 JavaScript 渲染頁面")
    paras = [(i + 1, p) for i, p in enumerate(re.split(r"\n{2,}", text))]
    segments = _pack_paragraphs(paras, lambda p: url)
    return {
        "video_id": _source_id(url),
        "source_type": "url",
        "title": title or url,
        "url": url,
        "segments": segments,
    }


def ingest_url_tavily(url: str) -> dict:
    if not os.environ.get("TAVILY_API_KEY"):
        raise ValueError("缺少 TAVILY_API_KEY")
    _assert_public_url(url)  # 同樣防護：不讓 server 被當成掃內網的跳板（經 Tavily 代抓）
    import requests

    resp = requests.post(
        "https://api.tavily.com/extract",
        json={"api_key": os.environ["TAVILY_API_KEY"], "urls": [url]},
        timeout=60,
    )
    resp.raise_for_status()
    results = resp.json().get("results", [])
    if not results or not results[0].get("raw_content"):
        raise ValueError("Tavily 抽取失敗")
    text = results[0]["raw_content"]
    paras = [(i + 1, p) for i, p in enumerate(re.split(r"\n{2,}", text))]
    segments = _pack_paragraphs(paras, lambda p: url)
    return {
        "video_id": _source_id(url),
        "source_type": "url",
        "title": url,
        "url": url,
        "segments": segments,
    }


def fetch_source_data(source: str | Path, engine: Optional[str] = None, custom_title: Optional[str] = None) -> dict:
    if isinstance(source, Path) or (isinstance(source, str) and not source.startswith(("http://", "https://"))):
        p = Path(source)
        if not p.exists():
            raise FileNotFoundError(f"找不到檔案: {p}")
        suffix = p.suffix.lower()
        if suffix == ".pdf":
            result = ingest_pdf_llamaparse(p) if engine == "llamaparse" else ingest_pdf_local(p)
        elif suffix == ".docx":
            result = ingest_docx(p)
        elif suffix == ".doc":
            raise ValueError("舊版 .doc 格式無法解析，請先用 Word 另存為 .docx")
        else:
            raise ValueError(f"不支援的檔案格式: {suffix} (支援 .pdf, .docx)")
        if custom_title:
            result["title"] = custom_title
        return result

    url_str = str(source)
    if YT_PAT.search(url_str):
        result = ingest_youtube(url_str)
    elif url_str.startswith(("http://", "https://")):
        result = ingest_url_tavily(url_str) if engine == "tavily" else ingest_url_free(url_str)
    else:
        raise ValueError(f"無法辨識的來源: {source}")

    if custom_title:
        result["title"] = custom_title
    return result


# ---------------- 向量庫入庫 ----------------

def aggregate_segments(segments: list[dict], window: int = WINDOW_SECONDS) -> list[dict]:
    blocks: list[dict] = []
    cur: Optional[dict] = None
    for seg in segments:
        if cur is None:
            cur = {"start": seg["start"], "end": seg["end"], "text": seg["text"]}
            continue
        if seg["end"] - cur["start"] <= window:
            cur["text"] += " " + seg["text"]
            cur["end"] = seg["end"]
        else:
            blocks.append(cur)
            cur = {"start": seg["start"], "end": seg["end"], "text": seg["text"]}
    if cur:
        blocks.append(cur)
    return blocks


def to_documents(source_data: dict) -> list[Document]:
    src_type = source_data.get("source_type", "unknown")
    if src_type == "youtube":
        blocks = aggregate_segments(source_data["segments"])
        mk_ref = lambda start: f"https://youtu.be/{source_data['video_id']}?t={int(start)}"
    else:
        blocks = source_data["segments"]
        mk_ref = lambda start: next(
            (s.get("ref") for s in source_data["segments"] if s["start"] == start),
            source_data.get("url", "")
        )

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        separators=["\n\n", "\n", "。", "！", "？", " ", ""],
    )
    docs = []
    chunk_index = 0
    for blk in blocks:
        chunks = splitter.split_text(blk["text"])
        for c in chunks:
            ref_url = mk_ref(blk["start"])
            docs.append(Document(
                page_content=c,
                metadata={
                    "video_id": source_data["video_id"],
                    "title": source_data.get("title", ""),
                    "start": blk["start"],
                    "end": blk["end"],
                    "url_at_time": ref_url,
                    "chunk_index": chunk_index,
                    "source_type": src_type,
                }
            ))
            chunk_index += 1
    return docs


def ingest_to_vectordb(
    source_data: dict,
    persist_dir: str = "./chroma_db",
    collection_name: str = "yt_rag",
    clear_existing: bool = True,
) -> list[Document]:
    docs = to_documents(source_data)
    if not docs:
        raise ValueError("切塊後沒有產出任何 Document")

    embeddings = GoogleGenerativeAIEmbeddings(model=EMBED_MODEL)
    db = Chroma(
        collection_name=collection_name,
        embedding_function=embeddings,
        persist_directory=persist_dir,
    )
    if clear_existing:
        try:
            db.delete(where={"video_id": source_data["video_id"]})
        except Exception:
            pass

    db.add_documents(docs)
    return docs


# ---------------- 知識圖譜抽取 & Neo4j 寫入 ----------------

def extract_triples(
    chunks: list[dict],
    model: str = CHAT_MODEL,
    progress_callback: Optional[Callable[[int, int, int], None]] = None,
) -> list[dict]:
    llm = ChatGoogleGenerativeAI(model=model, temperature=0)
    all_triples = []
    total = len(chunks)
    for i, chunk in enumerate(chunks):
        try:
            resp = llm.invoke(EXTRACT_PROMPT + chunk["text"])
            raw = resp.text.strip().removeprefix("```json").removeprefix("```").removesuffix("```")
            data = json.loads(raw)
            for t in data.get("triples", []):
                if all(k in t and t[k] for k in ("subject", "relation", "object")):
                    t["chunk_id"] = chunk["chunk_id"]
                    all_triples.append(t)
        except Exception as e:
            print(f"[!] 抽取 chunk {i} 時出錯: {e}")
        if progress_callback:
            progress_callback(i + 1, total, len(all_triples))
    return all_triples


def write_to_neo4j(triples: list[dict], chunks: list[dict], video: dict):
    driver = GraphDatabase.driver(
        os.environ.get("NEO4J_URI", "bolt://localhost:7687"),
        auth=(os.environ.get("NEO4J_USER", "neo4j"), os.environ["NEO4J_PASSWORD"]),
    )
    with driver.session() as s:
        # 冪等清空同來源舊圖
        s.run("MATCH (c:Chunk {video_id: $vid}) DETACH DELETE c", vid=video["video_id"])
        s.run(
            "MERGE (v:Video {id: $vid}) SET v.title = $title, v.url = $url, v.source_type = $st",
            vid=video["video_id"],
            title=video.get("title", ""),
            url=video.get("url", ""),
            st=video.get("source_type", "unknown"),
        )
        for c in chunks:
            s.run(
                """MERGE (ck:Chunk {id: $cid}) SET ck.text = $text, ck.start = $start,
                        ck.video_id = $vid, ck.url_at_time = $url
                   WITH ck MATCH (v:Video {id: $vid}) MERGE (ck)-[:PART_OF]->(v)""",
                cid=c["chunk_id"],
                text=c["text"],
                start=c["start"],
                vid=video["video_id"],
                url=c["url_at_time"],
            )
        for t in triples:
            s.run(
                """MERGE (a:Entity {name: $subj})
                   MERGE (b:Entity {name: $obj})
                   MERGE (a)-[r:REL {type: $rel}]->(b)
                   WITH a, b MATCH (ck:Chunk {id: $cid})
                   MERGE (a)-[:MENTIONED_IN]->(ck)
                   MERGE (b)-[:MENTIONED_IN]->(ck)""",
                subj=t["subject"].strip(),
                obj=t["object"].strip(),
                rel=t["relation"].strip(),
                cid=t["chunk_id"],
            )
    driver.close()


def run_full_pipeline(
    source: str | Path,
    engine: Optional[str] = None,
    custom_title: Optional[str] = None,
    progress_callback: Optional[Callable[[str, int, int], None]] = None,
) -> dict:
    """
    完整端到端執行：
    1. 擷取來源 (抓取字幕/解析PDF/解析DOCX/網頁)
    2. 切塊與向量入庫 (ChromaDB)
    3. LLM 抽取知識圖譜與入庫 (Neo4j)
    """
    if progress_callback:
        progress_callback("parsing", 0, 100)

    # 1. 擷取來源
    source_data = fetch_source_data(source, engine=engine, custom_title=custom_title)

    if progress_callback:
        progress_callback("vectorizing", 25, 100)

    # 2. 向量入庫
    docs = ingest_to_vectordb(source_data)

    if progress_callback:
        progress_callback("graphing", 50, 100)

    # 3. 轉成 chunk 字典準備抽取圖譜
    chunk_dicts = [
        {
            "chunk_id": f"{source_data['video_id']}_{i}",
            "text": doc.page_content,
            "start": doc.metadata["start"],
            "url_at_time": doc.metadata["url_at_time"],
        }
        for i, doc in enumerate(docs)
    ]

    def _graph_prog(cur, total, triples_so_far):
        if progress_callback:
            pct = 50 + int((cur / max(total, 1)) * 45)
            progress_callback(f"graphing_chunk_{cur}_{total}", pct, 100)

    triples = extract_triples(chunk_dicts, progress_callback=_graph_prog)

    # 4. 寫入 Neo4j
    write_to_neo4j(triples, chunk_dicts, source_data)

    if progress_callback:
        progress_callback("completed", 100, 100)

    return {
        "status": "success",
        "video_id": source_data["video_id"],
        "source_type": source_data["source_type"],
        "title": source_data["title"],
        "url": source_data["url"],
        "chunks_count": len(docs),
        "triples_count": len(triples),
    }
